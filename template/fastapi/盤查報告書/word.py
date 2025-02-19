from docx import Document

def create_word_file(filename: str):
    """ 產生與提供的Word文件格式一致的文件 """
    doc = Document()
    
    # 標題
    doc.add_heading('亞東科技大學板橋校區 溫室氣體盤查報告書', level=1)
    doc.add_paragraph('\n2025 年 01 月 15 日')
    
    # 目錄
    doc.add_heading('目錄', level=2)
    sections = [
        ('第一章、本校簡介與政策聲明', 2),
        ('1.1 前言', 2),
        ('1.2 學校簡介', 2),
        ('1.3 組織及架構', 3),
        ('1.4 報告書涵蓋期間與責任/有效期間', 4),
        ('1.5 宣告本盤查報告書製作之依據', 4),
        ('1.6 本盤查報告書製作目的', 4),
        ('第二章、盤查邊界設定', 5),
        ('2.1 組織邊界設定', 5),
        ('2.2 報告邊界', 5)
    ]
    for title, page in sections:
        doc.add_paragraph(f"{title} \t {page}")
    
    # 內容示例
    doc.add_heading('第一章、本校簡介與政策聲明', level=2)
    doc.add_heading('前言', level=3)
    doc.add_paragraph(
        '本校創校迄今，歷任校長遵循創辦人創校職志，經營擘畫，積極發揚「誠、勤、樸、慎、創新」精神形成優良校風...'
    )
    
    doc.add_heading('學校簡介', level=3)
    doc.add_paragraph(
        '亞東科技大學於民國五十七年十月，在遠東集團創辦人徐有庠先生的「弘文明德，育才興國」理念下創設...'
    )
    
    # 表格示例
    doc.add_heading('表1、學校場所資料表', level=3)
    table = doc.add_table(rows=5, cols=2)
    table.cell(0, 0).text = '學校名稱'
    table.cell(0, 1).text = '亞東科技大學'
    table.cell(1, 0).text = '校長'
    table.cell(1, 1).text = '黃茂全'
    table.cell(2, 0).text = '教職員生總人數'
    table.cell(2, 1).text = '4,397人'
    table.cell(3, 0).text = '學校地址'
    table.cell(3, 1).text = '新北市板橋區四川路二段58號'
    
    doc.save(filename)
