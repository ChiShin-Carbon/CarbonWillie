from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches
from docx.shared import Cm

def set_ch5_table1(table, electricity_usage=None):
    header_cells = [
        (0, 0, "全廠電力"),
        (1, 0, "全校電力\n(仟度)"), (1, 1, "全校火力電力\n(仟度)"),(1, 2,"風力\n(仟度)"),(1, 3, "水力\n(仟度)"), (1, 4, "地熱\n(仟度)"), (1, 5,"潮汐\n(仟度)"), (1, 6,"其他再生能源\n(仟度)"),
        (1, 7,  "其他再生能源\n備註"), (1, 8, "核能發電量\n(仟度)"), (1, 9, "其他發電量\n(仟度)"), (1, 10, "其他發電量\n備註"),
        (1, 11, "全廠蒸氣產生量\n(仟度)")
    ]

    # 添加API數據行
    data_cells = [
        (2, 0, str(electricity_usage) if electricity_usage is not None else "-"), 
        (2, 1, "-"), (2, 2, "-"), (2, 3, "-"), (2, 4, "-"), 
        (2, 5, "-"), (2, 6, "-"), (2, 7, "-"), (2, 8, "-"), 
        (2, 9, "-"), (2, 10, "-"), (2, 11, "-")
    ]

    # 合併標題和數據單元格
    all_cells = header_cells + data_cells

    for row_idx, col_idx, text in all_cells:
        cell = table.cell(row_idx, col_idx)
        paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        run = paragraph.add_run(text)

    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._element.get_or_add_tcPr()
            tc_borders = OxmlElement('w:tcBorders')

            for border_name in ['w:top', 'w:left', 'w:bottom', 'w:right', 'w:insideH', 'w:insideV']:
                border = OxmlElement(border_name)
                border.set(qn('w:val'), 'single')  # 單線
                border.set(qn('w:sz'), '4')       # 線條大小（1/8 pt）
                border.set(qn('w:space'), '0')    # 無間距
                border.set(qn('w:color'), '000000')  # 黑色
                tc_borders.append(border)

            tc_pr.append(tc_borders)

    table.cell(0, 0).merge(table.cell(0, 11))  # 合併第1列

    for cell in table.rows[0].cells:  # 只針對第一列
        shading = parse_xml(r'<w:shd {} w:fill="FFF2CC"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading)  # 設置背景顏色

        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 讓標題水平置中
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # 這句是為了強制將文字置中

        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        run.font.bold = True  # 設置為粗體

         # 垂直置中
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER  # 垂直置中

    # 設置表格內字體
    for row in table.rows:
        for cell in row.cells:
            paragraph = cell.paragraphs[0]
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()

            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(0)
            paragraph_format.space_after = Pt(0)
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

            # 設定字型
            font = run.font
            font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "標楷體")
            font.size = Pt(9)

             # 設定水平置中
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  

            # 設定垂直置中
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER



def set_ch5_table2(table, quantitative_inventory=None):
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.shared import Pt
    
    # Default values in case quantitative_inventory is None
    if quantitative_inventory is None:
        quantitative_inventory = {
            "total_emission_equivalent": 0,
            "CO2_emission_equivalent": 0,
            "CH4_emission_equivalent": 0,
            "N2O_emission_equivalent": 0,
            "HFCS_emission_equivalent": 0,
            "PFCS_emission_equivalent": 0,
            "SF6_emission_equivalent": 0,
            "NF3_emission_equivalent": 0
        }
    
    # Extract emission values
    co2 = quantitative_inventory.get("CO2_emission_equivalent", 0)
    ch4 = quantitative_inventory.get("CH4_emission_equivalent", 0)
    n2o = quantitative_inventory.get("N2O_emission_equivalent", 0)
    hfcs = quantitative_inventory.get("HFCS_emission_equivalent", 0)
    pfcs = quantitative_inventory.get("PFCS_emission_equivalent", 0)
    sf6 = quantitative_inventory.get("SF6_emission_equivalent", 0)
    nf3 = quantitative_inventory.get("NF3_emission_equivalent", 0)
    total = quantitative_inventory.get("total_emission_equivalent", 0)
    
    # Calculate percentages
    def calculate_percentage(value, total):
        if total == 0:
            return "0.00%"
        elif value == 0:
            return "-----"
        else:
            return f"{(value / total * 100):.2f}%"
    
    co2_percent = calculate_percentage(co2, total)
    ch4_percent = calculate_percentage(ch4, total)
    n2o_percent = calculate_percentage(n2o, total)
    hfcs_percent = calculate_percentage(hfcs, total)
    pfcs_percent = calculate_percentage(pfcs, total)
    sf6_percent = calculate_percentage(sf6, total)
    nf3_percent = calculate_percentage(nf3, total)
    
    # Format emission values to 3 decimal places
    co2_str = f"{co2:.3f}" if co2 != 0 else "0.000"
    ch4_str = f"{ch4:.3f}" if ch4 != 0 else "0.000"
    n2o_str = f"{n2o:.3f}" if n2o != 0 else "0.000"
    hfcs_str = f"{hfcs:.3f}" if hfcs != 0 else "0.000"
    pfcs_str = f"{pfcs:.3f}" if pfcs != 0 else "0.000"
    sf6_str = f"{sf6:.3f}" if sf6 != 0 else "0.000"
    nf3_str = f"{nf3:.3f}" if nf3 != 0 else "0.000"
    total_str = f"{total:.3f}" if total != 0 else "0.000"
    
    header_cells = [
        (0, 0, "全廠七大溫室氣體排放量統計表"),
        (1, 0, "溫室氣體"), (1, 1, "CO2"), (1, 2, "CH4"), (1, 3, "N2O"), (1, 4, "HFCs"), 
        (1, 5, "PFCs"), (1, 6, "SF6"), (1, 7, "NF3"), (1, 8, "年總排放當量"), (1, 9, "生質排放當量"),
        (2, 0, "排放當量\n(公噸CO2e/年)"),
        (2, 1, co2_str), (2, 2, ch4_str), (2, 3, n2o_str), (2, 4, hfcs_str),
        (2, 5, pfcs_str), (2, 6, sf6_str), (2, 7, nf3_str), (2, 8, total_str), (2, 9, "0.000"),
        (3, 0, "氣體別占比\n(％)"),
        (3, 1, "100.00%"), (3, 2, ch4_percent), (3, 3, n2o_percent), (3, 4, hfcs_percent),
        (3, 5, pfcs_percent), (3, 6, sf6_percent), (3, 7, nf3_percent), (3, 8, "100.00%"), (3, 9, "-----"),
        (4, 0, "註：依溫室氣體排放量盤查登錄管理辦法第二條第一款規定，溫室氣體排放量以公噸二氧化碳當量(公噸CO2e)表示，並四捨五入至小數點後第三位。")
    ]

    for row_idx, col_idx, text in header_cells:
        cell = table.cell(row_idx, col_idx)
        paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        run = paragraph.add_run(text)

    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._element.get_or_add_tcPr()
            tc_borders = OxmlElement('w:tcBorders')

            for border_name in ['w:top', 'w:left', 'w:bottom', 'w:right', 'w:insideH', 'w:insideV']:
                border = OxmlElement(border_name)
                border.set(qn('w:val'), 'single')  # 單線
                border.set(qn('w:sz'), '4')       # 線條大小（1/8 pt）
                border.set(qn('w:space'), '0')    # 無間距
                border.set(qn('w:color'), '000000')  # 黑色
                tc_borders.append(border)

            tc_pr.append(tc_borders)

    table.cell(0, 0).merge(table.cell(0, 9))  # 合併第1列
    table.cell(4, 0).merge(table.cell(4, 9))  # 合併最後一列

    for cell in table.rows[0].cells:  # 只針對第一列
        shading = parse_xml(r'<w:shd {} w:fill="FFF2CC"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading)  # 設置背景顏色

        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 讓標題水平置中
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # 這句是為了強制將文字置中

        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        run.font.bold = True  # 設置為粗體
           
    # 設置表格內字體 & 文字對齊
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):  
            paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()

            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(0)
            paragraph_format.space_after = Pt(0)
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

            # 設定字型
            font = run.font
            font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "標楷體")
            font.size = Pt(9)

            # 設定對齊方式
            if row_idx == 4:  # 註解部分
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 置左
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 其他部分保持置中

            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER  # 上下置中


def set_ch5_table3(table, quantitative_inventory=None):
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.shared import Pt
    
    # Default values in case quantitative_inventory is None
    if quantitative_inventory is None:
        quantitative_inventory = {
            "category1_total_emission_equivalent": 0,
            "category1_CO2_emission_equivalent": 0,
            "category1_CH4_emission_equivalent": 0,
            "category1_N2O_emission_equivalent": 0,
            "category1_HFCS_emission_equivalent": 0,
            "category1_PFCS_emission_equivalent": 0,
            "category1_SF6_emission_equivalent": 0,
            "category1_NF3_emission_equivalent": 0
        }
    
    # Extract category 1 emission values
    co2 = quantitative_inventory.get("category1_CO2_emission_equivalent", 0)
    ch4 = quantitative_inventory.get("category1_CH4_emission_equivalent", 0)
    n2o = quantitative_inventory.get("category1_N2O_emission_equivalent", 0)
    hfcs = quantitative_inventory.get("category1_HFCS_emission_equivalent", 0)
    pfcs = quantitative_inventory.get("category1_PFCS_emission_equivalent", 0)
    sf6 = quantitative_inventory.get("category1_SF6_emission_equivalent", 0)
    nf3 = quantitative_inventory.get("category1_NF3_emission_equivalent", 0)
    total = quantitative_inventory.get("category1_total_emission_equivalent", 0)
    
    # Calculate percentages
    def calculate_percentage(value, total):
        if total == 0:
            return "0.00%"
        elif value == 0:
            return "-----"
        else:
            return f"{(value / total * 100):.2f}%"
    
    co2_percent = calculate_percentage(co2, total)
    ch4_percent = calculate_percentage(ch4, total)
    n2o_percent = calculate_percentage(n2o, total)
    hfcs_percent = calculate_percentage(hfcs, total)
    pfcs_percent = calculate_percentage(pfcs, total)
    sf6_percent = calculate_percentage(sf6, total)
    nf3_percent = calculate_percentage(nf3, total)
    
    # Format emission values to 3 decimal places
    co2_str = f"{co2:.3f}" if co2 != 0 else "0.000"
    ch4_str = f"{ch4:.3f}" if ch4 != 0 else "0.000"
    n2o_str = f"{n2o:.3f}" if n2o != 0 else "0.000"
    hfcs_str = f"{hfcs:.3f}" if hfcs != 0 else "0.000"
    pfcs_str = f"{pfcs:.3f}" if pfcs != 0 else "0.000"
    sf6_str = f"{sf6:.3f}" if sf6 != 0 else "0.000"
    nf3_str = f"{nf3:.3f}" if nf3 != 0 else "0.000"
    total_str = f"{total:.3f}" if total != 0 else "0.000"
    
    header_cells = [
        (0, 0, "類別一、七大溫室氣體排放量統計表"),
        (1, 0, "溫室氣體"), (1, 1, "CO2"), (1, 2, "CH4"), (1, 3, "N2O"), (1, 4, "HFCs"), 
        (1, 5, "PFCs"), (1, 6, "SF6"), (1, 7, "NF3"), (1, 8, "年總排放當量"), (1, 9, "生質排放當量"),
        (2, 0, "排放當量\n(公噸CO2e/年)"),
        (2, 1, co2_str), (2, 2, ch4_str), (2, 3, n2o_str), (2, 4, hfcs_str),
        (2, 5, pfcs_str), (2, 6, sf6_str), (2, 7, nf3_str), (2, 8, total_str), (2, 9, "0.000"),
        (3, 0, "氣體別占比\n(％)"),
        (3, 1, "100.00%"), (3, 2, ch4_percent), (3, 3, n2o_percent), (3, 4, hfcs_percent),
        (3, 5, pfcs_percent), (3, 6, sf6_percent), (3, 7, nf3_percent), (3, 8, "100.00%"), (3, 9, "-----")
    ]

    for row_idx, col_idx, text in header_cells:
        cell = table.cell(row_idx, col_idx)
        paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        run = paragraph.add_run(text)

    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._element.get_or_add_tcPr()
            tc_borders = OxmlElement('w:tcBorders')

            for border_name in ['w:top', 'w:left', 'w:bottom', 'w:right', 'w:insideH', 'w:insideV']:
                border = OxmlElement(border_name)
                border.set(qn('w:val'), 'single')  # 單線
                border.set(qn('w:sz'), '4')       # 線條大小（1/8 pt）
                border.set(qn('w:space'), '0')    # 無間距
                border.set(qn('w:color'), '000000')  # 黑色
                tc_borders.append(border)

            tc_pr.append(tc_borders)

    table.cell(0, 0).merge(table.cell(0, 9))  # 合併第1列

    for cell in table.rows[0].cells:  # 只針對第一列
        shading = parse_xml(r'<w:shd {} w:fill="FFF2CC"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading)  # 設置背景顏色

        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 讓標題水平置中
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # 這句是為了強制將文字置中

        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        run.font.bold = True  # 設置為粗體

        # 垂直置中
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER  # 垂直置中
           
    # 設置表格內字體
    for row in table.rows:
        for cell in row.cells:
            paragraph = cell.paragraphs[0]
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()

            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(0)
            paragraph_format.space_after = Pt(0)
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

            # 設定字型
            font = run.font
            font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "標楷體")
            font.size = Pt(9)

            # 設定水平置中
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  
            
            # 設定垂直置中
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER



def set_ch5_table4(table):
    header_cells = [
        (0, 0, "全廠溫室氣體範疇別及類別一與二排放型式排放量統計表"),
        (1, 0, "範疇"),(1, 1, "類別一"),(1, 5, "類別二"),(1, 6, "總排放當量"),
        (2, 1, "固定排放"),(2, 2, "製程排放"),(2, 3, "移動排放"),(2, 4, "逸散排放"),(2, 5, "外購電力"),
        (3, 0, "排放當量\n(公噸CO2e/年)"),(5, 0, "氣體別占比\n(%)"),
        (7, 0, "註：依溫室氣體排放量盤查登錄管理辦法第二條第一款規定，溫室氣體排放量以公噸二氧化碳當量(公噸CO2e)表示，並四捨五入至小數點後第三位。"),
    ]

    for row_idx, col_idx, text in header_cells:
        cell = table.cell(row_idx, col_idx)
        paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        run = paragraph.add_run(text)

    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._element.get_or_add_tcPr()
            tc_borders = OxmlElement('w:tcBorders')

            for border_name in ['w:top', 'w:left', 'w:bottom', 'w:right', 'w:insideH', 'w:insideV']:
                border = OxmlElement(border_name)
                border.set(qn('w:val'), 'single')  # 單線
                border.set(qn('w:sz'), '4')       # 線條大小（1/8 pt）
                border.set(qn('w:space'), '0')    # 無間距
                border.set(qn('w:color'), '000000')  # 黑色
                tc_borders.append(border)

            tc_pr.append(tc_borders)

    # 進行合併
    table.cell(0, 0).merge(table.cell(0, 6)) 
    table.cell(7, 0).merge(table.cell(7, 6))  # 註解部分
    table.cell(1, 1).merge(table.cell(1, 4))  
    table.cell(3, 1).merge(table.cell(3, 4))  
    table.cell(5, 1).merge(table.cell(5, 4))

    table.cell(1, 0).merge(table.cell(2, 0))  
    table.cell(3, 0).merge(table.cell(4, 0))
    table.cell(5, 0).merge(table.cell(6, 0))  

    table.cell(1, 6).merge(table.cell(2, 6))  
    table.cell(3, 6).merge(table.cell(4, 6))
    table.cell(5, 6).merge(table.cell(6, 6)) 
    

    for cell in table.rows[0].cells:  # 只針對第一列
        shading = parse_xml(r'<w:shd {} w:fill="FFF2CC"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading)  # 設置背景顏色

        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 讓標題水平置中

        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        run.font.bold = True  # 設置為粗體

    # 設置表格內字體 & 文字對齊
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()

            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(0)
            paragraph_format.space_after = Pt(0)
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

            # 設定字型
            font = run.font
            font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "標楷體")
            font.size = Pt(9)

            # 設定對齊方式
            if row_idx == 7:  # 註解部分
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 置左
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 其他部分保持置中

            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER  # 上下置中

                # 設定特定儲存格 (1,1) 和 (1,5) 為粗體
            if (row_idx == 1 and col_idx in [1, 5]):
                run.font.bold = True  # 設置為粗體

