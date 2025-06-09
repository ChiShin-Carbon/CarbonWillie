from docx import Document
from docx.shared import Cm

from .storeDef2 import get_org_name

from .ch0Def import set_heading, set_heading2, set_paragraph, set_explain
from .ch3Def import set_ch3_table1,set_ch3_table2
from .storeDef3 import get_latest_baseline_year
from . storeDef4 import get_total_emission_equivalent


def create_chapter3(user_id):
    org_name = get_org_name(user_id)
    year= get_latest_baseline_year()
    total=get_total_emission_equivalent()

    doc = Document()

        # 獲取文檔的第一個 section（默認只有一個）
    section = doc.sections[0]

    # 設置自訂邊界，單位是厘米 (cm)
    section.top_margin = Cm(2)  # 上邊距
    section.bottom_margin = Cm(2)  # 下邊距
    section.left_margin = Cm(2)  # 左邊距
    section.right_margin = Cm(2)  # 右邊距
        
    ################第三章######################
    #標題
    title = doc.add_heading("第三章、報告溫室氣體排放量",level=1)
    set_heading(title)
    
    #3.1 溫室氣體排放類型與排放量說明
    preface = doc.add_heading("3.1 溫室氣體排放類型與排放量說明",level=2)
    set_heading2(preface)
    
    content = doc.add_paragraph("經盤查，本機構排放之溫室氣體種類主要有二氧化碳(CO2)、氧化亞氮(N2O)、甲烷(CH4)及氫氟碳化物(HFCs)四類。其中，二氧化碳(CO2)排放主要來自【消防設施（滅火器）、其他發電引擎（緊急發電機）及外購電力】，甲烷(CH4)的排放來自【化糞池及其他發電引擎(緊急發電機)】，氧化亞氮(N2O) 排放來自【其他發電引擎（緊急發電機）】，氫氟碳化物(HFCs)的排放來自廠區內【消防設施（滅火器）、各式冰水機（冰水主機）、飲水機及冷氣機】的冷媒逸散。")
    set_paragraph(content)

    #3.2 直接溫室氣體排放（類別1排放）
    preface = doc.add_heading("3.2 直接溫室氣體排放（類別1排放）",level=2)
    set_heading2(preface)
    
    content = doc.add_paragraph("本機構直接溫室氣體排放源，如表3-1所示。")
    set_paragraph(content)

    table_explain = doc.add_paragraph(f"表3-1、{org_name}直接溫室氣體排放源")
    set_explain(table_explain)
    # 新增表格
    table1 = doc.add_table(rows=4, cols=16)
    
    set_ch3_table1(table1)


    #3.3 能源間接溫室氣體排放（類別2排放）
    preface = doc.add_heading("3.3 能源間接溫室氣體排放（類別2排放）",level=2)
    set_heading2(preface)
    
    content = doc.add_paragraph("本機構能源間接溫室氣體排放源，如表3-2所示。")
    set_paragraph(content)

    table_explain = doc.add_paragraph(f"表3-2、{org_name}能源間接溫室氣體排放源")
    set_explain(table_explain)
    
    table2 = doc.add_table(rows=3, cols=16)
    set_ch3_table2(table2)



    
    #3.4 溫室氣體總排放量
    preface = doc.add_heading("3.4 溫室氣體總排放量",level=2)
    set_heading2(preface)
    
    content = doc.add_paragraph(f"經盤查，本機構{year}年溫室氣體總排放量{total}公噸CO2e。")
    set_paragraph(content)


    # 插入分頁符號
    doc.add_page_break()

    
    return doc
