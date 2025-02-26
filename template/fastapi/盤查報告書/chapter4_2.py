from docx import Document
from docx.shared import Cm

from storeDef import set_heading, set_heading2, set_paragraph, set_explain,set_ch4_table1,set_ch4_table2,set_ch4_table3,set_ch4_stairs1,set_ch4_stairs2,set_ch4_stairs3,set_ch4_stairs4,set_ch4_stairs5,set_ch4_stairs6

def add_stairs_paragraphs(doc, texts, format_func):
    for text in texts:
        para = doc.add_paragraph(text)
        format_func(para) 

def create_chapter4_2():
    doc = Document()

        # 獲取文檔的第一個 section（默認只有一個）
    section = doc.sections[0]

    # 設置自訂邊界，單位是厘米 (cm)
    section.top_margin = Cm(2)  # 上邊距
    section.bottom_margin = Cm(2)  # 下邊距
    section.left_margin = Cm(2)  # 左邊距
    section.right_margin = Cm(2)  # 右邊距
  
################################################################################
    stairs2 = doc.add_paragraph("D. 逸散排放源（冷媒）：")
    set_ch4_stairs2(stairs2)

    stairs3 = doc.add_paragraph("(A) 溫室氣體排放量計算公式如下：")
    set_ch4_stairs3(stairs3)

    stairs4 = doc.add_paragraph("a. 溫室氣體排放量 = 溫室氣體逸散量 × 全球暖化潛勢值(GWP)")
    set_ch4_stairs4(stairs4)

    stairs4 = doc.add_paragraph("b. （當年未有添加紀錄）\n冷媒逸散量量化方式 = 冷媒原始填充量 × 設備逸散率(%)")
    set_ch4_stairs4(stairs4)

    stairs4 = doc.add_paragraph("c. （當年有添加紀錄）冷媒逸散量量化方式=實際填充量")
    set_ch4_stairs4(stairs4)

    stairs3 = doc.add_paragraph("(B) 冷媒原始填充量(ton)。")
    set_ch4_stairs3(stairs3)

    stairs3 = doc.add_paragraph("(C) 依IPCC建議值（冷媒逸散率排放因子），並取中間值計算，如表4-10所示。")
    set_ch4_stairs3(stairs3)

    table_explain = doc.add_paragraph("表4-6、設備之冷媒逸散率排放因子")
    set_explain(table_explain)
    table1 = doc.add_table(rows=9, cols=3)
    set_ch4_table2(table1)

    table_explain = doc.add_paragraph("表4-7、逸散排放源（冷媒）排放源HFCs")
    set_explain(table_explain)
    table2 = doc.add_table(rows=15, cols=16)
    set_ch4_table1(table2)

    doc.add_paragraph("")  # 插入一個空白行
############################################################################
    stairs2 = doc.add_paragraph("E.	製程排放：機構內並無製程紀錄，本項次無對應活動數據，故無對應之盤查結果可供揭露。")
    set_ch4_stairs2(stairs2)

    stairs3 = doc.add_paragraph("(A) 溫室氣體排放量計算公式如下：\n溫室氣體排放量 = 活動數據 × 排放係數 × 全球暖化潛勢值(GWP)")
    set_ch4_stairs3(stairs3)

    stairs3 = doc.add_paragraph("(B) 活動數據：盤查年份的購置數量（公噸）")
    set_ch4_stairs3(stairs3)

    stairs3 = doc.add_paragraph("(C) 排放係數：生產過程所造成的溫室氣體排放。量化方法採用質能平衡法，以下舉常用的乙炔、焊條為例。")
    set_ch4_stairs3(stairs3)

    stairs4_texts = [
    "乙炔燃燒排放（氣焊）：",
    "· 活動數據：盤查年份的購置數量（公斤）",
    "· C2H2 + 2.5 O2 -> 2CO2 + H2O",
    "· 每燃燒1 mole C2H2（分子量26），產生2 mole CO2（分子量88）",
    "焊條燃燒排放（電焊）：",
    "· 活動數據：盤查年份，購置數量（公斤） ×焊條含碳率(%)",
    "· C + O2 -> CO2",
    "· 每燃燒1 mole C（分子量12），產生1 mole CO2（分子量44）",
    "· CO2排放係數 = 44/12 = 3.667 公噸/公噸C"
    ]

    add_stairs_paragraphs(doc, stairs4_texts, set_ch4_stairs4)

    table_explain = doc.add_paragraph("表4-8、製程排放源排放源CO2s")
    set_explain(table_explain)
    table3 = doc.add_table(rows=3, cols=16)
    set_ch4_table1(table3)

    doc.add_paragraph("")  # 插入一個空白行
    ############################################################################

    stairs1 = doc.add_paragraph("(2) 類別2 – 能源間接排放")
    set_ch4_stairs1(stairs1)

    stairs2 = doc.add_paragraph("A.	間接排放源（外購電力）:")
    set_ch4_stairs2(stairs2)

    stairs3 = doc.add_paragraph("(A) 溫室氣體排放量計算公式如下：\n溫室氣體排放量 = 活動數據 × 排放係數 × 全球暖化潛勢值(GWP)")
    set_ch4_stairs3(stairs3)

    stairs3 = doc.add_paragraph("(B) 活動數據：全年用電量（千度）")
    set_ch4_stairs3(stairs3)

    stairs3 = doc.add_paragraph("(C) 排放係數：113年度之電力排碳係數為0.495公斤CO2e/度")
    set_ch4_stairs3(stairs3)

    table_explain = doc.add_paragraph("表4-9、間接排放源（外購電力）排放源")
    set_explain(table_explain)
    table4 = doc.add_table(rows=3, cols=16)
    set_ch4_table1(table4)




    stairs5 = doc.add_paragraph("4.1.1 活動數據蒐集與轉換方式")
    set_ch4_stairs5(stairs5)
    stairs6 = doc.add_paragraph("(1) 本機構各排放源之量化公式與活動數據蒐集方式彙整如表4-10所示。")
    set_ch4_stairs5(stairs6)
    stairs6 = doc.add_paragraph("(2) 各種溫室氣體之排放依來源不同，將活動數據單位化為公噸、公秉、千度等單位。")
    set_ch4_stairs5(stairs6)

    table_explain = doc.add_paragraph("表4-10、活動數據蒐集彙整表")
    set_explain(table_explain)
    table4 = doc.add_table(rows=6, cols=5)
    table4.cell(0, 0).text = "營運邊界"
    table4.cell(0, 1).text = "量化方式"
    table4.cell(0, 2).text = "排放源"
    table4.cell(0, 3).text = "負責部門"
    table4.cell(0, 4).text = "活動數據收集說明"

    table4.cell(1, 0).text = "直接排放源"
    table4.cell(1, 1).text = "排放係數法"
    table4.cell(2, 1).text = "排放係數法"
    table4.cell(3, 1).text = "估算溫室氣體\n逸散量"
    table4.cell(4, 1).text = "估算溫室氣體\n逸散量"
    table4.cell(5, 0).text = "能源間接\n排放源"
    table4.cell(5, 1).text = "排放係數法"

    table4.cell(1, 2).text = "化糞池"
    table4.cell(2, 2).text = "消防活動（滅火器）"
    table4.cell(3, 2).text = "冷媒補充－各式冰水機、飲水機、冷氣機"
    table4.cell(4, 2).text = "緊急發電機（柴油）"
    table4.cell(5, 2).text = "外購電力"

    table4.cell(1, 4).text = "人事考勤系統"
    table4.cell(2, 4).text = "消防設備調查表（滅火器）"
    table4.cell(3, 4).text = "冷媒銘牌填充量"
    table4.cell(4, 4).text = "採購單據"
    table4.cell(5, 4).text = "亞東科技大學板橋校區台電電費單\n(電號：nn-nn-nnnn-nn-n)"

    set_ch4_table3(table4)



    stairs5 = doc.add_paragraph("4.1.2 排放係數來源")
    set_ch4_stairs5(stairs5)
    content = doc.add_paragraph("針對各種不同的溫室氣體排放源，本次盤查採用之排放係數來源主要為「溫室氣體排放係數管理表6.0.4版」，部分排放係數參考IPCC AR6；本次盤查採用溫室氣體盤查登錄表3.0.0文件，請詳見附件二。")
    set_paragraph(content)
    
    stairs5 = doc.add_paragraph("4.1.3 全球暖化潛勢值(GWP)")
    set_ch4_stairs5(stairs5)
    content = doc.add_paragraph("計算出各類溫室氣體排放量後，應乘上各種溫室氣體所屬之全球暖化潛勢值(GWP)，並將其計算結果轉化為CO2e，單位為公噸/年。")
    set_paragraph(content)




    return doc
