from docx import Document
from docx.shared import Cm

from .storeDef import set_heading, set_heading2, set_paragraph, set_ch4_stairs7

def add_stairs_paragraphs(doc, texts, format_func):
    for text in texts:
        para = doc.add_paragraph(text)
        format_func(para) 

def create_chapter4_3(user_id):
    doc = Document()

        # 獲取文檔的第一個 section（默認只有一個）
    section = doc.sections[0]

    # 設置自訂邊界，單位是厘米 (cm)
    section.top_margin = Cm(2)  # 上邊距
    section.bottom_margin = Cm(2)  # 下邊距
    section.left_margin = Cm(2)  # 左邊距
    section.right_margin = Cm(2)  # 右邊距
  
################################################################################
     #4.2 量化方法變更說明
    preface = doc.add_heading("4.2 量化方法變更說明",level=2)
    set_heading2(preface)
    
    content = doc.add_paragraph("量化方法改變時，本機構除以新的量化計算方式計算外，並需與原來之計算方式做一比較，並說明二者之差異及選用新方法的理由。")
    set_paragraph(content)

  #4.3 排放係數與變更說明
    preface = doc.add_heading("4.3 排放係數與變更說明",level=2)
    set_heading2(preface)
    
    content = doc.add_paragraph("本次盤查作業若量化方法屬於排放係數法者。")
    set_paragraph(content)

    #4.4 有效位數
    preface = doc.add_heading("4.4 有效位數",level=2)
    set_heading2(preface)
    
    content = doc.add_paragraph("有關本機構溫室氣體盤查作業之有效位數設定，係參考環境部公告「國家溫室氣體登錄平台運算方式第5版」之建議進行，採四捨五入取到小數點後三位。")
    set_paragraph(content)

    #4.5 重大排放源之資訊流
    preface = doc.add_heading("4.5	重大排放源之資訊流",level=2)
    set_heading2(preface)
    
    content = doc.add_paragraph("根據本機構進行的溫室氣體盤查結果，總排放量為【xxxx.xxxx】公噸。其中，範疇二的外購電力排放量為【xxxx.xxxx】公噸，占總排放量的【xx.xx】%。相比之下，範疇一的排放量為【xxxx.xxxx】公噸，占總排放量的【xx.xx】%，顯示出外購電力在整體溫室氣體排放量中占據了最大的比例。各式活動源及各類溫室氣體的排放量及排放占比，請詳見表5.1。")
    set_paragraph(content)

    #4.6 本次盤查排除事項、注意事項及推估說明
    preface = doc.add_heading("4.6 本次盤查排除事項、注意事項及推估說明",level=2)
    set_heading2(preface)
    
    content = doc.add_paragraph("以下為本次辦理溫室氣體盤查工作，有關盤查排除事項、注意事項及活動數據不完整資訊下推估的說明。")
    set_paragraph(content)

    stairs7 = doc.add_paragraph("． 本機構消防設備於有痒科技大樓設有ABC型乾粉滅火器257支、誠勤大樓設有ABC型乾粉滅火器149支、元智大樓設有ABC型乾粉滅火器148支、實習大樓設有ABC型乾粉滅火器108支、亞東第一停車場設有ABC型乾粉滅火器77支，因該式滅火器並不會產生溫室氣體，故而將其排除不計。*清冊中未標示乾粉滅火器型式*")
    set_ch4_stairs7(stairs7)
    stairs7 = doc.add_paragraph("． 本機構消防設備另設有T10型BC乾粉滅火器17支，皆於2015年購入，經查明2024年皆未有使用以及填充紀錄；基於完整性原則，本次盤查全數計入排放量。*尚未確認*")
    set_ch4_stairs7(stairs7)
    stairs7 = doc.add_paragraph("． 本機構消防設備另設有FM200海龍滅火器1支，於2016年購入，經查明2024年未有使用以及填充紀錄；基於完整性原則，本次盤查計入排放量。*尚未確認*")
    set_ch4_stairs7(stairs7)
    stairs7 = doc.add_paragraph("． 本機構緊急發電機，經與管理單位協議並確認，參照最近兩次添購備用柴油紀錄以及目測油箱庫存量，推算取得2024年的年度使用量。*清冊中尚無緊急發電機之資訊*")
    set_ch4_stairs7(stairs7)
    stairs7 = doc.add_paragraph("． 本機構汙水下水道工程雖已竣工，但經查證(桃園市下水道雲端智慧管理系統，https://sewergis.tycg.gov.tw/Account/Login?ReturnUrl=%2F)現階段尚未正式通水，因此仍以校區員工年度總工時作為化糞池逸散排放的計算依據。*尚未確認*")
    set_ch4_stairs7(stairs7)
    stairs7 = doc.add_paragraph("． 本機構未有乙炔、焊條及其他與製程相關油品或氣體使用紀錄。本機構無半導體製程，故無全氟碳化物(PFCs)、六氟化硫(SF6)及三氟化氮(NF3)氣體逸散。")
    set_ch4_stairs7(stairs7)
    stairs7 = doc.add_paragraph("． 其他間接排放（類別三、類別五及類別六），包括成品委外運輸、員工上下班及商務出差、自動販賣機等其它間接排放，因無法掌控其活動及溫室氣體排放，2024年度只進行排放源鑑別之工作，不予以量化。*尚未確認*")
    set_ch4_stairs7(stairs7)




    # 插入分頁符號
    doc.add_page_break()


    return doc
