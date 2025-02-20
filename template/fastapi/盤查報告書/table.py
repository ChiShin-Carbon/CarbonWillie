from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_table_borders(table):
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._element.get_or_add_tcPr()
            tc_borders = OxmlElement('w:tcBorders')
            
            for border_name in ['w:top', 'w:left', 'w:bottom', 'w:right', 'w:insideH', 'w:insideV']:
                border = OxmlElement(border_name)
                border.set(qn('w:val'), 'single')  # 單線
                border.set(qn('w:sz'), '4')       # 線條大小（1/8 pt）
                border.set(qn('w:space'), '0')    # 無間距
                border.set(qn('w:color'), '000000')  # 黑色
                tc_borders.append(border)
            
            tc_pr.append(tc_borders)

doc = Document()
table = doc.add_table(rows=2, cols=3)
table.style = 'Table Grid'  # 設定表格樣式

# 設定框線
set_table_borders(table)

# 插入一些測試數據
for i, row in enumerate(table.rows):
    for j, cell in enumerate(row.cells):
        cell.text = f'R{i+1}C{j+1}'
        cell.paragraphs[0].runs[0].font.size = Pt(12)

doc.save("table.docx")
