from docx import Document
from docx.shared import Cm


from .storeDef2 import get_org_name,get_charge_person,get_org_address,get_intro,get_summary

from .ch0Def import set_heading, set_heading2, set_paragraph, set_explain
from .ch1Def import set_ch1_table1,set_ch1_pointlist


def create_chapter1(user_id):
    org_name = get_org_name(user_id)
    org_address = get_org_address(user_id)
    charge_person = get_charge_person(user_id)
    intro=get_intro(user_id)
    summary = get_summary(user_id)

    doc = Document()

    # 獲取文檔的第一個 section（默認只有一個）
    section = doc.sections[0]

    # 設置自訂邊界，單位是厘米 (cm)
    section.top_margin = Cm(2)  # 上邊距
    section.bottom_margin = Cm(2)  # 下邊距
    section.left_margin = Cm(2)  # 左邊距
    section.right_margin = Cm(2)  # 右邊距
    
    ################第一章######################
    #標題
    title = doc.add_heading("第一章、機構簡介與政策聲明",level=1)
    set_heading(title)
    
    #1.1前言
    preface = doc.add_heading("1.1 前言",level=2)
    set_heading2(preface)
    
    content = doc.add_paragraph(f"{intro}")
    set_paragraph(content)

    #1.2簡介
    preface = doc.add_heading("1.2 簡介",level=2)
    set_heading2(preface)
    
    content = doc.add_paragraph(f"{summary}")
    set_paragraph(content)

    table_explain = doc.add_paragraph("表1、機構場所資料表")
    set_explain(table_explain)
    table = doc.add_table(rows=4, cols=2)
    table.cell(0, 0).text = "機構名稱"
    table.cell(0, 1).text = f"{org_name}"
    table.cell(1, 0).text = "負責人"
    table.cell(1, 1).text = f"{charge_person}"
    table.cell(2, 0).text = "員工總人數"
    table.cell(2, 1).text = "n人"
    table.cell(3, 0).text = "機構地址"
    table.cell(3, 1).text = f"{org_address}"
    set_ch1_table1(table)


    # 插入分頁符號
    doc.add_page_break()

    #1.3組織及架構
    preface = doc.add_heading("1.3 組織及架構",level=2)
    set_heading2(preface)
    content = doc.add_paragraph("【請插入組織架構圖】")
    set_paragraph(content)

    photo_explain = doc.add_paragraph(f"圖一、{org_name}組織架構圖")
    set_explain(photo_explain)

    # 插入分頁符號
    doc.add_page_break()

    
    #1.4	報告書涵蓋期間與責任/有效期間
    preface = doc.add_heading("1.4 報告書涵蓋期間與責任/有效期間",level=2)
    set_heading2(preface)
    content1_4_1 = doc.add_paragraph("1.4.1 報告書涵蓋期間與責任")
    content1_4_1_1=doc.add_paragraph(f"本報告書之盤查內容是以【OOOO年度】於{org_address}（以下均稱本機構）組織邊界範圍內產生之所有溫室氣體為盤查範圍，並供作下年度新報告書完成前引用。")
    
    content1_4_2 = doc.add_paragraph("1.4.2 本報告書為隔年1月時開始進行前一年度之溫室氣體排放量之各項盤查工作，並於2月開始報告書之內容製作，其涵蓋前一年本校之溫室氣體排放總結，供作本年度及下年度新報告書完成前引用。")
    content1_4_3=doc.add_paragraph("1.4.3 報告書完成後，經過年度內部諮詢之程序，並修正缺失後，完成本報告書。")
    content1_4_4=doc.add_paragraph("1.4.4 本報告書盤查範圍只限於本機構營運範圍之總溫室氣體之排放量，本機構之組織營運範圍，若有變動時，本報告書將一併進行修正並重新發行。")

    set_ch1_pointlist(content1_4_1)
    set_ch1_pointlist(content1_4_1_1)
    set_ch1_pointlist(content1_4_2)
    set_ch1_pointlist(content1_4_3)
    set_ch1_pointlist(content1_4_4)
    
    #1.5 宣告本盤查報告書製作之依據
    preface = doc.add_heading("1.5 宣告本盤查報告書製作之依據",level=2)
    set_heading2(preface)
    content1_5 = doc.add_paragraph("本報告書乃根據 ISO 14064-1：2018（CNS 14064-1：2022）進行盤查與計算。")
    set_paragraph(content1_5)


    #1.6 本盤查報告書製作目的
    preface = doc.add_heading("1.6 本盤查報告書製作目的",level=2)
    set_heading2(preface)
    content1_6_1 = doc.add_paragraph("1.6.1 展現本機構溫室氣體盤查結果。")
    content1_6_2=doc.add_paragraph("1.6.2 妥當紀錄本機構溫室氣體排放清冊，以利社會責任標準查證之需求。")

    set_ch1_pointlist(content1_6_1)
    set_ch1_pointlist(content1_6_2)


    # 插入分頁符號
    doc.add_page_break()



    return doc