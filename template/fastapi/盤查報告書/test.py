from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_heading(paragraph, text, level=1):
    """ 設定標題格式 """
    run = paragraph.add_run(text)
    run.font.name = 'Times New Roman'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')  # 設定中文字體
    run.font.size = Pt(18 if level == 1 else 16)
    run.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER if level == 1 else WD_PARAGRAPH_ALIGNMENT.LEFT
    
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph_format.space_before = Pt(18 if level == 1 else 12)
    paragraph_format.space_after = Pt(12)

def set_paragraph_format(paragraph):
    """ 設定段落格式 """
    run = paragraph.add_run()
    run.font.name = 'Times New Roman'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')  # 設定中文字體
    run.font.size = Pt(12)
    
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph_format.line_spacing = 1.15
    paragraph_format.space_after = Pt(6)
    

def create_word_file(filename: str):
    doc = Document()
    
    ######################### 第一章 #########################
    chapter1 = doc.add_paragraph()
    set_heading(chapter1, "第一章、本校簡介與政策聲明", level=1)
    
    # 前言
    heading2 = doc.add_paragraph()
    set_heading(heading2, "1.1 前言", level=2)
    ch1_para1 = doc.add_paragraph("本校創校迄今，歷任校長遵循創辦人創校職志，經營擘畫，積極發揚「誠、勤、樸、慎、創新」精神形成優良校風，並秉持「創意、務實、宏觀、合作、溝通、熱忱」的教育理念，以科技與人文融匯、創新與品質並重、專業與通識兼顧、理論與實務結合為主軸，發展為實務化、資訊化、人文化、創新化、國際化的高等學府。")
    ch1_para2 = doc.add_paragraph("為提供學生多元學習，整合相關學術資源，本校特成立電通、工程、醫護暨管理三大學院，藉由各學系的合作、因應產業需求，開設相關學程，讓學生透過跨領域學習，提升專業知能與職場競爭力。")
    ch1_para3 = doc.add_paragraph("本校積極提升教學、研究、輔導與服務外，並與遠傳、新世紀資通、遠東新世紀、亞東醫院等遠東集團產學合作，成果斐然，已成為技職教育新典範。")

    set_paragraph_format(ch1_para1)
    set_paragraph_format(ch1_para2)
    set_paragraph_format(ch1_para3)
    
    # 學校簡介
    heading2 = doc.add_paragraph()
    set_heading(heading2, "學校簡介", level=2)
    para2 = doc.add_paragraph("亞東科技大學於民國五十七年十月，在遠東集團創辦人徐有庠先生...亞東科技大學。")
    set_paragraph_format(para2)
    
    # 表格範例
    heading2 = doc.add_paragraph()
    set_heading(heading2, "表1、學校場所資料表", level=2)
    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    data = [
        ("學校名稱", "亞東科技大學"),
        ("校長", "黃茂全"),
        ("教職員生總人數", "4,397人"),
        ("學校地址", "新北市板橋區四川路二段58號")
    ]
    
    for i, (col1, col2) in enumerate(data):
        row = table.rows[i]
        row.cells[0].text = col1
        row.cells[1].text = col2
    
    # 插入組織架構圖（佔位）
    heading2 = doc.add_paragraph()
    set_heading(heading2, "【請放置組織架構圖】", level=2)
    doc.add_paragraph("圖一、亞東科技大學組織架構圖")
    
    # 第二章
    chapter2 = doc.add_paragraph()
    set_heading(chapter2, "第二章、盤查邊界設定", level=1)
    
    # 存檔
    doc.save(filename)
    print(f"檔案已儲存至: {filename}")

if __name__ == "__main__":
    output_filename = "溫室氣體盤查報告書.docx"
    create_word_file(output_filename)
