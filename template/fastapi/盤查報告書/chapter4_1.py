from docx import Document
from docx.shared import Cm

from .ch0Def import set_heading, set_heading2, set_paragraph, set_explain
from .ch4Def import set_ch4_table1,set_ch4_stairs1,set_ch4_stairs2,set_ch4_stairs3,set_ch4_stairs4,set_ch4_table1_1,set_ch4_table1_2,set_ch4_table1_3,set_ch4_table1_4,set_ch4_table1_5,set_ch4_table1_6,set_ch4_table1_7

def create_chapter4_1(user_id):
    doc = Document()

        # 獲取文檔的第一個 section（默認只有一個）
    section = doc.sections[0]

    # 設置自訂邊界，單位是厘米 (cm)
    section.top_margin = Cm(2)  # 上邊距
    section.bottom_margin = Cm(2)  # 下邊距
    section.left_margin = Cm(2)  # 左邊距
    section.right_margin = Cm(2)  # 右邊距

    def add_stairs_paragraphs(doc, texts, format_func):
        for text in texts:
            para = doc.add_paragraph(text)
            format_func(para) 
        
    ###############第四章######################
    #標題
    title = doc.add_heading("第四章、數據品質管理",level=1)
    set_heading(title)
    
    #4.1 量化方法
    preface = doc.add_heading("4.1 量化方法",level=2)
    set_heading2(preface)
    
    content = doc.add_paragraph("本機構各種溫室氣體排放量計算方式主要採用「排放係數法」計算。")
    set_paragraph(content)

    stairs1 = doc.add_paragraph("(1) 類別1 – 直接排放")
    set_ch4_stairs1(stairs1)
################################################################################

    stairs2 = doc.add_paragraph("A.	移動燃燒排放源（公務車）")
    set_ch4_stairs2(stairs2)

    stairs3 = doc.add_paragraph("(A) 溫室氣體排放量計算公式如下：\n溫室氣體年排放量 ＝ 車用汽油使用量 ×［CO2排放係數 × CO2溫暖化潛勢+ CH4排放係數 × CH4溫暖化潛勢 + N2O排放係數 × N2O溫暖化潛勢］× 車用汽油低位熱值")
    set_ch4_stairs3(stairs3)



    table_explain = doc.add_paragraph("表4-1、移動燃燒排放源（公務車）CO2")
    set_explain(table_explain)
    table11 = doc.add_table(rows=3, cols=16)
    set_ch4_table1_1(table11)

    table_explain = doc.add_paragraph("表4-2、移動燃燒排放源（公務車）CH4")
    set_explain(table_explain)
    table12 = doc.add_table(rows=3, cols=16)
    set_ch4_table1_2(table12)

    table_explain = doc.add_paragraph("表4-3、移動燃燒排放源（公務車）N2O")
    set_explain(table_explain)
    table13 = doc.add_table(rows=3, cols=16)
    set_ch4_table1_3(table13)

    doc.add_paragraph("")  # 插入一個空白行

    ############################################################################
    stairs2 = doc.add_paragraph("B. 固定燃燒排放源（緊急發電機）：")
    set_ch4_stairs2(stairs2)

    stairs3 = doc.add_paragraph("(A) 溫室氣體排放量計算公式如下：\n溫室氣體排放量 = 活動數據 × 排放係數 × 全球暖化潛勢值(GWP)")
    set_ch4_stairs3(stairs3)

    stairs3 = doc.add_paragraph("(B) 活動數據：汽油用量（公噸）、柴油用量（公秉）")
    set_ch4_stairs3(stairs3)

    stairs3 = doc.add_paragraph("(C) 排放係數：溫室氣體排放係數管理表6.0.4版。")
    set_ch4_stairs3(stairs3)

    table_explain = doc.add_paragraph("表4-4、固定燃燒排放源（緊急發電機）CO2")
    set_explain(table_explain)
    table1 = doc.add_table(rows=3, cols=16)
    set_ch4_table1_4(table1)

    table_explain = doc.add_paragraph("表4-5、固定燃燒排放源（緊急發電機）CH4")
    set_explain(table_explain)
    table2 = doc.add_table(rows=3, cols=16)
    set_ch4_table1_5(table2)

    table_explain = doc.add_paragraph("表4-6、固定燃燒排放源（緊急發電機）N2O")
    set_explain(table_explain)
    table3 = doc.add_table(rows=3, cols=16)
    set_ch4_table1_6(table3)

    doc.add_paragraph("")  # 插入一個空白行
############################################################################
    stairs2 = doc.add_paragraph("C.	逸散排放源（化糞池）：")
    set_ch4_stairs2(stairs2)

    stairs3 = doc.add_paragraph("(A) 溫室氣體排放量計算公式如下：\n溫室氣體排放量 = 活動數據 × 排放係數 × 全球暖化潛勢值(GWP)")
    set_ch4_stairs3(stairs3)

    stairs3 = doc.add_paragraph("(B) 活動數據：統計全年人小時")
    set_ch4_stairs3(stairs3)

    stairs3 = doc.add_paragraph("(C) 排放係數：溫室氣體排放係數管理表6.0.4版(6_逸散排放源)之化糞池係數，並換算為人時0.0000015938公噸/人時。")
    set_ch4_stairs3(stairs3)

    table_explain = doc.add_paragraph("表4-7、逸散排放源（化糞池）CH4")
    set_explain(table_explain)
    table4 = doc.add_table(rows=3, cols=16)
    set_ch4_table1(table4)
    
    doc.add_paragraph("")  # 插入一個空白行
    ############################################################################
    stairs2 = doc.add_paragraph("D.	逸散排放源（滅火器）：")
    set_ch4_stairs2(stairs2)

    stairs3 = doc.add_paragraph("(A) 溫室氣體排放量計算公式如下：\n溫室氣體排放量 = 活動數據 × 排放係數 × 全球暖化潛勢值(GWP)")
    set_ch4_stairs3(stairs3)

    stairs3 = doc.add_paragraph("(B) 活動數據：CO2滅火器該年度使用量（公噸）")
    set_ch4_stairs3(stairs3)

    stairs3 = doc.add_paragraph("(C) 排放係數：質量平衡係數CO2為1。")
    set_ch4_stairs3(stairs3)

    table_explain = doc.add_paragraph("表4-8、逸散排放源（滅火器）CO2、HFCS")
    set_explain(table_explain)
    table5 = doc.add_table(rows=3, cols=16)
    set_ch4_table1(table5)

    doc.add_paragraph("")  # 插入一個空白行

    return doc
