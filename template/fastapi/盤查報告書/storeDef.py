from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches
from docx.shared import Cm

def set_heading(paragraph):
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    font = run.font
    font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "標楷體")
    font.size = Pt(18)
    font.color.rgb = RGBColor(0, 0, 0)  # 設置字體為黑色
    
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(18)
    paragraph_format.space_after = Pt(12)

    paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 設定標題置中

def set_heading2(paragraph):
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    font = run.font
    font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "標楷體")
    font.size = Pt(16)
    font.color.rgb = RGBColor(0, 0, 0)  # 設置字體為黑色
    
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(12)
    paragraph_format.space_after = Pt(12)

    paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

def set_paragraph(paragraph):
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    font = run.font
    font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "標楷體")
    font.size = Pt(12)
    
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(6)

    paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph_format.line_spacing = 1.15

    # **設定首行縮排 (2 個中文字元寬度)**
    paragraph_format.first_line_indent = Pt(24)  # 一般中文字大小是 12pt，2 個字就是 24pt

def set_explain(paragraph):
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    font = run.font
    font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "標楷體")
    font.size = Pt(10)
    
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(12)
    paragraph_format.space_after = Pt(0)
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 設定標題置中

def set_ch1_table1(table):
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._element.get_or_add_tcPr()
            tc_borders = OxmlElement('w:tcBorders')
            
            for border_name in ['w:top', 'w:left', 'w:bottom', 'w:right', 'w:insideH', 'w:insideV']:
                border = OxmlElement(border_name)
                border.set(qn('w:val'), 'single')  # 單線
                border.set(qn('w:sz'), '4')       # 線條大小（1/8 pt）
                border.set(qn('w:space'), '0')    # 無間距
                border.set(qn('w:color'), '000000')  # 黑色
                tc_borders.append(border)
            
            tc_pr.append(tc_borders)

    for row in table.rows:
        cell = row.cells[0]  # 選取第一列的單元格
        shading = parse_xml(r'<w:shd {} w:fill="E2EFD9"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading)

    for row in table.rows:
            for cell in row.cells:
                paragraph = cell.paragraphs[0]
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()

                paragraph_format = paragraph.paragraph_format
                paragraph_format.space_before = Pt(0)
                paragraph_format.space_after = Pt(0)
                paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                        
                # 設定字型
                font = run.font
                font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn('w:eastAsia'), "標楷體")
                font.size = Pt(11)
                
                # 設定對齊方式
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def set_ch1_pointlist(paragraph):
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    font = run.font
    font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "標楷體")
    font.size = Pt(12)
    
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(6)

    paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph_format.line_spacing = 1.15

def set_ch2_table1(table):
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._element.get_or_add_tcPr()
            tc_borders = OxmlElement('w:tcBorders')

            for border_name in ['w:top', 'w:left', 'w:bottom', 'w:right', 'w:insideH', 'w:insideV']:
                border = OxmlElement(border_name)
                border.set(qn('w:val'), 'single')  # 單線
                border.set(qn('w:sz'), '4')       # 線條大小（1/8 pt）
                border.set(qn('w:space'), '0')    # 無間距
                border.set(qn('w:color'), '000000')  # 黑色
                tc_borders.append(border)

            tc_pr.append(tc_borders)
    
    # 設置第一列（Header Row）的背景色、粗體與置中
    for cell in table.rows[0].cells:  # 只針對第一列
        shading = parse_xml(r'<w:shd {} w:fill="E2EFD9"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading)  # 設置背景顏色

        paragraph = cell.paragraphs[0]
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        run.font.bold = True  # 設置為粗體
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 讓標題置中

    # 設置第一欄（Column 1）內容置中，第二欄（Column 2）內容靠左
    for row in table.rows[1:]:  # 跳過第一列，從第二列開始
        first_col = row.cells[0]  # 第一欄
        second_col = row.cells[1]  # 第二欄

        # 第一欄置中
        first_paragraph = first_col.paragraphs[0]
        first_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 第二欄靠左
        second_paragraph = second_col.paragraphs[0]
        second_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 設置表格內字體
    for row in table.rows:
        for cell in row.cells:
            paragraph = cell.paragraphs[0]
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()

            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(0)
            paragraph_format.space_after = Pt(0)
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

            # 設定字型
            font = run.font
            font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "標楷體")
            font.size = Pt(12)
    
     # 設定欄位寬度（第一欄30%，第二欄70%）
    table.autofit = False  # 取消自動調整大小，手動設定寬度
    for row in table.rows:
        row.cells[0].width = Inches(2)  # 第一欄 30%
        row.cells[1].width = Inches(5)  # 第二欄 70%

def set_ch3_table1(table):
    header_cells = [
        (0, 0, "製程名稱"), (0, 1, "設備名稱"), (0, 2, "原燃物料或產品"), (0, 5, "排放源資料"),
        (0, 7, "可能產生溫室氣體種類"),(0, 14, "是否屬汽電共生設備"),(0, 15, "備註*"),
        (1, 2, "類別"), (1, 3, "名稱"), (1, 4, "是否屬生質能源"), (1, 5,"範疇別"), (1, 6,"製程/逸散/外購電力類別"),
        (1, 7,  "CO2"), (1, 8, "CH4"), (1, 9, "N2O"), (1, 10, "HFCS"),
        (1, 11, "PFCS"), (1, 12, "SF6"), (1, 13,"NF3")
    ]

    for row_idx, col_idx, text in header_cells:
        cell = table.cell(row_idx, col_idx)
        paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        run = paragraph.add_run(text)

    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._element.get_or_add_tcPr()
            tc_borders = OxmlElement('w:tcBorders')

            for border_name in ['w:top', 'w:left', 'w:bottom', 'w:right', 'w:insideH', 'w:insideV']:
                border = OxmlElement(border_name)
                border.set(qn('w:val'), 'single')  # 單線
                border.set(qn('w:sz'), '4')       # 線條大小（1/8 pt）
                border.set(qn('w:space'), '0')    # 無間距
                border.set(qn('w:color'), '000000')  # 黑色
                tc_borders.append(border)

            tc_pr.append(tc_borders)

    # 合併第1, 2, 15, 16列的最上面兩行
    table.cell(0, 0).merge(table.cell(1, 0))  # 合併第1列
    table.cell(0, 1).merge(table.cell(1, 1))  # 合併第2列
    table.cell(0, 14).merge(table.cell(1, 14))  # 合併第15列
    table.cell(0, 15).merge(table.cell(1, 15))  # 合併第16列

    # 合併第3, 4, 5列的第一行為一個大格
    table.cell(0, 2).merge(table.cell(0, 4))  # 合併第3, 4, 5列
    table.cell(0, 5).merge(table.cell(0, 6))  # 合併第3, 4, 5列
    table.cell(0, 7).merge(table.cell(0, 13))  # 合併第3, 4, 5列

    
    
    # 設置第一列和第二列（Header Row）的背景色、置中
    for row_index in range(2):  # 迭代第一行和第二行
        for cell in table.rows[row_index].cells:  # 針對每一行的儲存格
            shading = parse_xml(r'<w:shd {} w:fill="CCFFFF"/>'.format(nsdecls('w')))
            cell._tc.get_or_add_tcPr().append(shading)  # 設置背景顏色

            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 讓標題水平置中
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # 這句是為了強制將文字置中

            # 垂直置中
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER  # 垂直置中
            
    # 設置從第3行開始的顏色
    for row_index in range(2, len(table.rows)):  # 從第3行開始
        # 設置第1, 2, 4列顏色為DBDBDB
        for col_index in [0, 1, 3]:
            shading = parse_xml(r'<w:shd {} w:fill="DBDBDB"/>'.format(nsdecls('w')))
            table.cell(row_index, col_index)._tc.get_or_add_tcPr().append(shading)

        # 設置第3, 5-15列顏色為F8CBAD
        for col_index in [2, 4, 5, 6, 7, 8, 9, 10, 11, 12, 13, 14]:
            shading = parse_xml(r'<w:shd {} w:fill="F8CBAD"/>'.format(nsdecls('w')))
            table.cell(row_index, col_index)._tc.get_or_add_tcPr().append(shading)


    # 設置表格內字體
    for row in table.rows:
        for cell in row.cells:
            paragraph = cell.paragraphs[0]
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()

            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(0)
            paragraph_format.space_after = Pt(0)
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

            # 設定字型
            font = run.font
            font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "標楷體")
            font.size = Pt(6)

def set_ch4_table1(table):
    
    header_cells = [
        (0, 0, "製程\n代碼"), (0, 1, "設備\n代碼"), (0, 2, "原燃物料或產品名稱"), (0, 3, "排放源資料"),
        (0, 5, "活動數據"), (0, 7, "排放係數(公噸/公噸or公秉or立方公尺)數據"),
        (1, 3, "範疇別"), (1, 4, "排放型式"), (1, 5, "活動\n數據"), (1, 6, "單位"),
        (1, 7, "溫室\n氣體"), (1, 8, "係數\n類型"), (1, 9, "預設排放係數"), (1, 10, "預設係數來源"),
        (1, 11, "係數\n單位"), (1, 12, "係數\n種類"), (1, 13, "排放量\n(公噸/年)"),
        (1, 14, "GWP"), (1, 15, "排放當量\n(公噸 CO2e/年)")
    ]

    for row_idx, col_idx, text in header_cells:
        cell = table.cell(row_idx, col_idx)
        paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        run = paragraph.add_run(text)

    for row in table.rows:
            for cell in row.cells:
                tc_pr = cell._element.get_or_add_tcPr()
                tc_borders = OxmlElement('w:tcBorders')

                for border_name in ['w:top', 'w:left', 'w:bottom', 'w:right', 'w:insideH', 'w:insideV']:
                    border = OxmlElement(border_name)
                    border.set(qn('w:val'), 'single')  # 單線
                    border.set(qn('w:sz'), '4')       # 線條大小（1/8 pt）
                    border.set(qn('w:space'), '0')    # 無間距
                    border.set(qn('w:color'), '000000')  # 黑色
                    tc_borders.append(border)

                tc_pr.append(tc_borders)

    # 合併第1, 2, 15, 16列的最上面兩行
    table.cell(0, 0).merge(table.cell(1, 0))  # 合併第1列
    table.cell(0, 1).merge(table.cell(1, 1))  # 合併第2列
    table.cell(0, 2).merge(table.cell(1, 2))  # 合併第3列

    table.cell(0, 3).merge(table.cell(0, 4))  
    table.cell(0, 5).merge(table.cell(0, 6))  
    table.cell(0, 7).merge(table.cell(0, 15))
    
    
    # 設置第一列和第二列（Header Row）的背景色、置中
    for row_index in range(2):  # 迭代第一行和第二行
        for cell in table.rows[row_index].cells:  # 針對每一行的儲存格
            shading = parse_xml(r'<w:shd {} w:fill="BDD6EE"/>'.format(nsdecls('w')))
            cell._tc.get_or_add_tcPr().append(shading)  # 設置背景顏色

            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 讓標題水平置中
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # 這句是為了強制將文字置中

            # 垂直置中
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER  # 垂直置中
            

    # 設置表格內字體
    for row in table.rows:
        for cell in row.cells:
            paragraph = cell.paragraphs[0]
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()

            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(0)
            paragraph_format.space_after = Pt(0)
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

            # 設定字型
            font = run.font
            font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "標楷體")
            font.size = Pt(7)
    
def set_ch4_table2(table):
    header_cells = [
        (0, 0, "設備名稱"), (0, 1, "排放因子(%)"), (0, 2, "防治設備回收率(%)"),
        (1, 0, "家用冷凍、冷藏裝備"), (1, 1, "0.1≦x≦0.5"), (1, 2, "70"), 
        (2, 0, "獨立商用冷凍、冷藏裝備"), (2, 1, "1≦x≦15"), (2, 2, "70"), 
        (3, 0, "中、大型冷凍、冷藏裝備"), (3, 1, "10≦x≦35"), (3, 2, "70"), 
        (4, 0, "交通用冷凍、冷藏裝備"), (4, 1, "15≦x≦50"), (4, 2, "70"), 
        (5, 0, "工業冷凍、冷藏裝備，包括食品加工及冷藏"), (5, 1, "7≦x≦25"), (5, 2, "90"), 
        (6, 0, "冰水機"), (6, 1, "2≦x≦15"), (6, 2, "95"), 
        (7, 0, "住宅及商業建築冷氣機"), (7, 1, "1≦x≦10"), (7, 2, "80"), 
        (8, 0, "移動式空氣清靜機"), (8, 1, "10≦x≦20"), (8, 2, "50"), 
    ]

    for row_idx, col_idx, text in header_cells:
        cell = table.cell(row_idx, col_idx)
        paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        run = paragraph.add_run(text)

    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._element.get_or_add_tcPr()
            tc_borders = OxmlElement('w:tcBorders')
            
            for border_name in ['w:top', 'w:left', 'w:bottom', 'w:right', 'w:insideH', 'w:insideV']:
                border = OxmlElement(border_name)
                border.set(qn('w:val'), 'single')  # 單線
                border.set(qn('w:sz'), '4')       # 線條大小（1/8 pt）
                border.set(qn('w:space'), '0')    # 無間距
                border.set(qn('w:color'), '000000')  # 黑色
                tc_borders.append(border)
            
            tc_pr.append(tc_borders)

    for cell in table.rows[0].cells:  # 只針對第一列
        shading = parse_xml(r'<w:shd {} w:fill="FBE4D5"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading)  # 設置背景顏色

    for row in table.rows:
            for cell in row.cells:
                paragraph = cell.paragraphs[0]
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()

                paragraph_format = paragraph.paragraph_format
                paragraph_format.space_before = Pt(0)
                paragraph_format.space_after = Pt(0)
                paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                        
                # 設定字型
                font = run.font
                font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn('w:eastAsia'), "標楷體")
                font.size = Pt(12)
                
                # 設定對齊方式
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def set_ch4_table3(table):

    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._element.get_or_add_tcPr()
            tc_borders = OxmlElement('w:tcBorders')
            
            for border_name in ['w:top', 'w:left', 'w:bottom', 'w:right', 'w:insideH', 'w:insideV']:
                border = OxmlElement(border_name)
                border.set(qn('w:val'), 'single')  # 單線
                border.set(qn('w:sz'), '4')       # 線條大小（1/8 pt）
                border.set(qn('w:space'), '0')    # 無間距
                border.set(qn('w:color'), '000000')  # 黑色
                tc_borders.append(border)
            
            tc_pr.append(tc_borders)

            # 設定垂直置中
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    table.cell(1, 0).merge(table.cell(4, 0))  # 合併第2列

    for cell in table.rows[0].cells:  # 只針對第一列
        shading = parse_xml(r'<w:shd {} w:fill="E2EFD9"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading)  # 設置背景顏色

    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            paragraph = cell.paragraphs[0]
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()

            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(0)
            paragraph_format.space_after = Pt(0)
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

            # 設定字型
            font = run.font
            font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "標楷體")
            font.size = Pt(10)

            # 設定對齊方式
            if row_idx == 0:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 第一列維持置中
            else:
                if col_idx in [0, 1]:  # 第一、二欄
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:  # 其他欄
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                

    # 設定欄位寬度
    table.autofit = False  # 取消自動調整大小，手動設定寬度
    total_width = Inches(7)  # 表格總寬度，可根據頁面大小調整
    column_widths = [0.12, 0.16, 0.30, 0.12, 0.30]  # 各欄比例
    for row in table.rows:
        for col_idx, width_ratio in enumerate(column_widths):
            row.cells[col_idx].width = total_width * width_ratio


def set_ch4_stairs1(paragraph):
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    font = run.font
    font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "標楷體")
    font.size = Pt(12)
    
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(6)

    paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph_format.line_spacing = 1.15

    paragraph_format.left_indent = Cm(1)  # 縮排 1 公分

def set_ch4_stairs2(paragraph):
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    font = run.font
    font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "標楷體")
    font.size = Pt(12)
    
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(6)

    paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph_format.line_spacing = 1.15

    paragraph_format.left_indent = Cm(1.75)  # 縮排 1 公分

def set_ch4_stairs3(paragraph):
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    font = run.font
    font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "標楷體")
    font.size = Pt(12)
    
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(6)

    paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph_format.line_spacing = 1.15

    paragraph_format.left_indent = Cm(2.25)  # 縮排 1 公分

def set_ch4_stairs4(paragraph):
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    font = run.font
    font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "標楷體")
    font.size = Pt(12)
    
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(6)

    paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph_format.line_spacing = 1.15

    paragraph_format.left_indent = Cm(3)  # 縮排 1 公分

def set_ch4_stairs5(paragraph):
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    font = run.font
    font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "標楷體")
    font.size = Pt(12)
    
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(12)
    paragraph_format.space_after = Pt(6)

    paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph_format.line_spacing = 1.15

    paragraph_format.left_indent = Cm(0.75)  # 縮排 1 公分

def set_ch4_stairs6(paragraph):
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    font = run.font
    font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "標楷體")
    font.size = Pt(12)
    
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(6)

    paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph_format.line_spacing = 1.15

    paragraph_format.left_indent = Cm(2.25)  # 縮排 1 公分


def set_ch4_stairs7(paragraph):
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    font = run.font
    font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "標楷體")
    font.size = Pt(12)
    
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(6)

    paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph_format.line_spacing = 1.15

    paragraph_format.left_indent = Cm(0.31)  # 縮排 1 公分
