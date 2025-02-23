from docx import Document
from docx.shared import Cm

from storeDef import set_heading, set_heading2, set_paragraph, set_explain,set_ch4_table1,set_ch4_stairs1,set_ch4_stairs2,set_ch4_stairs3,set_ch4_stairs4


def create_chapter4():
    doc = Document()

        # 獲取文檔的第一個 section（默認只有一個）
    section = doc.sections[0]

    # 設置自訂邊界，單位是厘米 (cm)
    section.top_margin = Cm(2)  # 上邊距
    section.bottom_margin = Cm(2)  # 下邊距
    section.left_margin = Cm(2)  # 左邊距
    section.right_margin = Cm(2)  # 右邊距
        
    ###############第四章######################
    #標題
    title = doc.add_heading("第四章、數據品質管理",level=1)
    set_heading(title)
    
    #4.1 量化方法
    preface = doc.add_heading("4.1 量化方法",level=2)
    set_heading2(preface)
    
    content = doc.add_paragraph("本公司各種溫室氣體排放量計算方式主要採用「排放係數法」計算。")
    set_paragraph(content)

    stairs1 = doc.add_paragraph("(1) 類別1 – 直接排放")
    set_ch4_stairs1(stairs1)

    stairs2 = doc.add_paragraph("A. 固定燃燒排放源（緊急發電機）：")
    set_ch4_stairs2(stairs2)

    stairs3 = doc.add_paragraph("(A) 溫室氣體排放量計算公式如下：\n溫室氣體排放量 = 活動數據 × 排放係數 × 全球暖化潛勢值(GWP)")
    set_ch4_stairs3(stairs3)

    stairs3 = doc.add_paragraph("(B) 活動數據：汽油用量（公噸）、柴油用量（公秉）")
    set_ch4_stairs3(stairs3)

    stairs3 = doc.add_paragraph("(C) 排放係數：溫室氣體排放係數管理表6.0.4版。")
    set_ch4_stairs3(stairs3)


    table_explain = doc.add_paragraph("表3-1、亞東科技大學直接溫室氣體排放源")
    set_explain(table_explain)
    # 新增表格
    table1 = doc.add_table(rows=15, cols=16)
     # 設定表頭
    table1.cell(0, 0).text = "製程代碼"
    table1.cell(0, 1).text = "設備代碼"
    table1.cell(0, 2).text = "原燃物料或產品"
    table1.cell(0, 5).text = "排放源資料"
    table1.cell(0, 7).text = "可能產生溫室氣體種類"
    table1.cell(0, 14).text ="是否屬汽電共生設備"
    table1.cell(0, 15).text = "備註*"

    table1.cell(1, 2).text = "範疇別"
    table1.cell(1, 3).text = "排放型式"
    table1.cell(1, 4).text = "是否屬生質能源"
    table1.cell(1, 5).text = "範疇別"
    table1.cell(1, 6).text = "製程/逸散/外購電力類別"
    table1.cell(1, 7).text = "CO2"
    table1.cell(1, 8).text = "CH4"
    table1.cell(1, 9).text = "N2O"
    table1.cell(1, 10).text = "HFCS"
    table1.cell(1, 11).text = "PFCS"
    table1.cell(1, 12).text = "SF6"
    table1.cell(1, 13).text = "NF3"

    set_ch4_table1(table1)




    return doc
