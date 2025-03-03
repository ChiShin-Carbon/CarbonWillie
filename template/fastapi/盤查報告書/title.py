from docx import Document
from docx.shared import Cm

from .storeDef import set_title
from .storeDef2 import get_org_name

from datetime import datetime


def create_title(user_id):
    org_name = get_org_name(user_id)
    doc = Document()

    # 獲取文檔的第一個 section（默認只有一個）
    section = doc.sections[0]

    # 設置自訂邊界，單位是厘米 (cm)
    section.top_margin = Cm(2)  # 上邊距
    section.bottom_margin = Cm(2)  # 下邊距
    section.left_margin = Cm(2)  # 左邊距
    section.right_margin = Cm(2)  # 右邊距
    
    ###################################################
    doc.add_paragraph("")  # 插入一個空白行
    doc.add_paragraph("")  # 插入一個空白行
    doc.add_paragraph("")  # 插入一個空白行
    title=doc.add_paragraph(f"OOOO年{org_name}\n溫室氣體盤查報告書")
    set_title(title)
    doc.add_paragraph("")  # 插入一個空白行
    doc.add_paragraph("")  # 插入一個空白行
    doc.add_paragraph("")  # 插入一個空白行
    doc.add_paragraph("")  # 插入一個空白行
    doc.add_paragraph("")  # 插入一個空白行
    doc.add_paragraph("")  # 插入一個空白行
    doc.add_paragraph("")  # 插入一個空白行
    doc.add_paragraph("")  # 插入一個空白行
    doc.add_paragraph("")  # 插入一個空白行
    doc.add_paragraph("")  # 插入一個空白行
    doc.add_paragraph("")  # 插入一個空白行
    doc.add_paragraph("")  # 插入一個空白行
    doc.add_paragraph("")  # 插入一個空白行
    doc.add_paragraph("")  # 插入一個空白行
    doc.add_paragraph("")  # 插入一個空白行
    doc.add_paragraph("")  # 插入一個空白行  doc.add_paragraph("")  # 插入一個空白行
    doc.add_paragraph("")  # 插入一個空白行
    doc.add_paragraph("")  # 插入一個空白行
    doc.add_paragraph("")  # 插入一個空白行

    # 取得當前日期
    current_date = datetime.now().strftime("%Y 年 %m 月 %d 日")
    title = doc.add_paragraph(current_date)
    set_title(title)

    # 插入分頁符號
    doc.add_page_break()



    return doc