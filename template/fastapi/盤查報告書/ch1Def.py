from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches
from docx.shared import Cm

def set_ch1_table1(table):
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._element.get_or_add_tcPr()
            tc_borders = OxmlElement('w:tcBorders')
            
            for border_name in ['w:top', 'w:left', 'w:bottom', 'w:right', 'w:insideH', 'w:insideV']:
                border = OxmlElement(border_name)
                border.set(qn('w:val'), 'single')  # 單線
                border.set(qn('w:sz'), '4')       # 線條大小（1/8 pt）
                border.set(qn('w:space'), '0')    # 無間距
                border.set(qn('w:color'), '000000')  # 黑色
                tc_borders.append(border)
            
            tc_pr.append(tc_borders)

    for row in table.rows:
        cell = row.cells[0]  # 選取第一列的單元格
        shading = parse_xml(r'<w:shd {} w:fill="E2EFD9"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading)

    for row in table.rows:
            for cell in row.cells:
                paragraph = cell.paragraphs[0]
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()

                paragraph_format = paragraph.paragraph_format
                paragraph_format.space_before = Pt(0)
                paragraph_format.space_after = Pt(0)
                paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                        
                # 設定字型
                font = run.font
                font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn('w:eastAsia'), "標楷體")
                font.size = Pt(11)
                
                # 設定對齊方式
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def set_ch1_pointlist(paragraph):
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    font = run.font
    font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "標楷體")
    font.size = Pt(12)
    
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(6)

    paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph_format.line_spacing = 1.15

