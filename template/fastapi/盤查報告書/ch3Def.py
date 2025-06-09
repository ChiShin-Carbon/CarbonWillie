from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches
from docx.shared import Cm


def set_ch3_table1(table):
    header_cells = [
        (0, 0, "製程名稱"), (0, 1, "設備名稱"), (0, 2, "原燃物料或產品"), (0, 5, "排放源資料"),
        (0, 7, "可能產生溫室氣體種類"),(0, 14, "是否屬汽電共生設備"),(0, 15, "備註*"),
        (1, 2, "類別"), (1, 3, "名稱"), (1, 4, "是否屬生質能源"), (1, 5,"範疇別"), (1, 6,"製程/逸散/外購電力類別"),
        (1, 7,  "CO2"), (1, 8, "CH4"), (1, 9, "N2O"), (1, 10, "HFCS"),
        (1, 11, "PFCS"), (1, 12, "SF6"), (1, 13,"NF3")
    ]

    # 設定資料值
    data_cells = [
        # 第3行資料 (row 2)
        (2, 0, "交通運輸工具"), (2, 1, "汽油引擎"), (2, 2, "1.原燃物料"), (2, 3, "車用汽油"),
        (2, 4, "否"), (2, 5, "範疇1"), (2, 6, ""), (2, 7, "v"), (2, 8, "v"), (2, 9, "v"),
        (2, 10, ""), (2, 11, ""), (2, 12, ""), (2, 13, ""), (2, 14, "否"), (2, 15, "公務車-汽油"),
        # 第4行資料 (row 3)
        (3, 0, "其他未分類製程"), (3, 1, "其他發電引擎"), (3, 2, "1.原燃物料"), (3, 3, "柴油"),
        (3, 4, "否"), (3, 5, "範疇1"), (3, 6, ""), (3, 7, "v"), (3, 8, "v"), (3, 9, "v"),
        (3, 10, ""), (3, 11, ""), (3, 12, ""), (3, 13, ""), (3, 14, "否"), (3, 15, "緊急發電機-柴油")
    ]

    # 填入標題
    for row_idx, col_idx, text in header_cells:
        cell = table.cell(row_idx, col_idx)
        paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        run = paragraph.add_run(text)

    # 填入資料
    for row_idx, col_idx, text in data_cells:
        cell = table.cell(row_idx, col_idx)
        paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        paragraph.clear()  # 清除現有內容
        run = paragraph.add_run(text)

    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._element.get_or_add_tcPr()
            tc_borders = OxmlElement('w:tcBorders')

            for border_name in ['w:top', 'w:left', 'w:bottom', 'w:right', 'w:insideH', 'w:insideV']:
                border = OxmlElement(border_name)
                border.set(qn('w:val'), 'single')  # 單線
                border.set(qn('w:sz'), '4')       # 線條大小（1/8 pt）
                border.set(qn('w:space'), '0')    # 無間距
                border.set(qn('w:color'), '000000')  # 黑色
                tc_borders.append(border)

            tc_pr.append(tc_borders)

    # 合併第1, 2, 15, 16列的最上面兩行
    table.cell(0, 0).merge(table.cell(1, 0))  # 合併第1列
    table.cell(0, 1).merge(table.cell(1, 1))  # 合併第2列
    table.cell(0, 14).merge(table.cell(1, 14))  # 合併第15列
    table.cell(0, 15).merge(table.cell(1, 15))  # 合併第16列

    # 合併第3, 4, 5列的第一行為一個大格
    table.cell(0, 2).merge(table.cell(0, 4))  # 合併第3, 4, 5列
    table.cell(0, 5).merge(table.cell(0, 6))  # 合併第3, 4, 5列
    table.cell(0, 7).merge(table.cell(0, 13))  # 合併第3, 4, 5列

    
    
    # 設置第一列和第二列（Header Row）的背景色、置中
    for row_index in range(2):  # 迭代第一行和第二行
        for cell in table.rows[row_index].cells:  # 針對每一行的儲存格
            shading = parse_xml(r'<w:shd {} w:fill="CCFFFF"/>'.format(nsdecls('w')))
            cell._tc.get_or_add_tcPr().append(shading)  # 設置背景顏色

            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 讓標題水平置中
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # 這句是為了強制將文字置中

            # 垂直置中
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER  # 垂直置中
            
    # 設置從第3行開始的顏色
    for row_index in range(2, len(table.rows)):  # 從第3行開始
        # 設置第1, 2, 4列顏色為DBDBDB
        for col_index in [0, 1, 3]:
            shading = parse_xml(r'<w:shd {} w:fill="DBDBDB"/>'.format(nsdecls('w')))
            table.cell(row_index, col_index)._tc.get_or_add_tcPr().append(shading)

        # 設置第3, 5-15列顏色為F8CBAD
        for col_index in [2, 4, 5, 6, 7, 8, 9, 10, 11, 12, 13, 14]:
            shading = parse_xml(r'<w:shd {} w:fill="F8CBAD"/>'.format(nsdecls('w')))
            table.cell(row_index, col_index)._tc.get_or_add_tcPr().append(shading)


    # 設置表格內字體
    for row in table.rows:
        for cell in row.cells:
            paragraph = cell.paragraphs[0]
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()

            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(0)
            paragraph_format.space_after = Pt(0)
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

            # 設定字型
            font = run.font
            font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "標楷體")
            font.size = Pt(6)



def set_ch3_table2(table):
    header_cells = [
        (0, 0, "製程名稱"), (0, 1, "設備名稱"), (0, 2, "原燃物料或產品"), (0, 5, "排放源資料"),
        (0, 7, "可能產生溫室氣體種類"),(0, 14, "是否屬汽電共生設備"),(0, 15, "備註*"),
        (1, 2, "類別"), (1, 3, "名稱"), (1, 4, "是否屬生質能源"), (1, 5,"範疇別"), (1, 6,"製程/逸散/外購電力類別"),
        (1, 7,  "CO2"), (1, 8, "CH4"), (1, 9, "N2O"), (1, 10, "HFCS"),
        (1, 11, "PFCS"), (1, 12, "SF6"), (1, 13,"NF3")
    ]

    # 設定資料值
    data_cells = [
        # 第3行資料 (row 2)
        (2, 0, "其他未分類製程"), (2, 1, "其他未歸類設施"), (2, 2, "1.原燃物料"), (2, 3, "其他電力"),
        (2, 4, "否"), (2, 5, "範疇2"), (2, 6, "併網"), (2, 7, "v"), (2, 8, ""), (2, 9, ""),
        (2, 10, ""), (2, 11, ""), (2, 12, ""), (2, 13, ""), (2, 14, "否"), (2, 15, "台電電力")
    ]

    # 填入標題
    for row_idx, col_idx, text in header_cells:
        cell = table.cell(row_idx, col_idx)
        paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        run = paragraph.add_run(text)

    # 填入資料
    for row_idx, col_idx, text in data_cells:
        cell = table.cell(row_idx, col_idx)
        paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        paragraph.clear()  # 清除現有內容
        run = paragraph.add_run(text)

    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._element.get_or_add_tcPr()
            tc_borders = OxmlElement('w:tcBorders')

            for border_name in ['w:top', 'w:left', 'w:bottom', 'w:right', 'w:insideH', 'w:insideV']:
                border = OxmlElement(border_name)
                border.set(qn('w:val'), 'single')  # 單線
                border.set(qn('w:sz'), '4')       # 線條大小（1/8 pt）
                border.set(qn('w:space'), '0')    # 無間距
                border.set(qn('w:color'), '000000')  # 黑色
                tc_borders.append(border)

            tc_pr.append(tc_borders)

    # 合併第1, 2, 15, 16列的最上面兩行
    table.cell(0, 0).merge(table.cell(1, 0))  # 合併第1列
    table.cell(0, 1).merge(table.cell(1, 1))  # 合併第2列
    table.cell(0, 14).merge(table.cell(1, 14))  # 合併第15列
    table.cell(0, 15).merge(table.cell(1, 15))  # 合併第16列

    # 合併第3, 4, 5列的第一行為一個大格
    table.cell(0, 2).merge(table.cell(0, 4))  # 合併第3, 4, 5列
    table.cell(0, 5).merge(table.cell(0, 6))  # 合併第3, 4, 5列
    table.cell(0, 7).merge(table.cell(0, 13))  # 合併第3, 4, 5列

    
    
    # 設置第一列和第二列（Header Row）的背景色、置中
    for row_index in range(2):  # 迭代第一行和第二行
        for cell in table.rows[row_index].cells:  # 針對每一行的儲存格
            shading = parse_xml(r'<w:shd {} w:fill="CCFFFF"/>'.format(nsdecls('w')))
            cell._tc.get_or_add_tcPr().append(shading)  # 設置背景顏色

            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 讓標題水平置中
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # 這句是為了強制將文字置中

            # 垂直置中
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER  # 垂直置中
            
    # 設置從第3行開始的顏色
    for row_index in range(2, len(table.rows)):  # 從第3行開始
        # 設置第1, 2, 4列顏色為DBDBDB
        for col_index in [0, 1, 3]:
            shading = parse_xml(r'<w:shd {} w:fill="DBDBDB"/>'.format(nsdecls('w')))
            table.cell(row_index, col_index)._tc.get_or_add_tcPr().append(shading)

        # 設置第3, 5-15列顏色為F8CBAD
        for col_index in [2, 4, 5, 6, 7, 8, 9, 10, 11, 12, 13, 14]:
            shading = parse_xml(r'<w:shd {} w:fill="F8CBAD"/>'.format(nsdecls('w')))
            table.cell(row_index, col_index)._tc.get_or_add_tcPr().append(shading)


    # 設置表格內字體
    for row in table.rows:
        for cell in row.cells:
            paragraph = cell.paragraphs[0]
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()

            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(0)
            paragraph_format.space_after = Pt(0)
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

            # 設定字型
            font = run.font
            font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "標楷體")
            font.size = Pt(6)