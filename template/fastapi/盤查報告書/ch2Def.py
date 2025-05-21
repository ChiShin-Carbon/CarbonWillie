from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches
from docx.shared import Cm

def set_ch2_table1(table):
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._element.get_or_add_tcPr()
            tc_borders = OxmlElement('w:tcBorders')

            for border_name in ['w:top', 'w:left', 'w:bottom', 'w:right', 'w:insideH', 'w:insideV']:
                border = OxmlElement(border_name)
                border.set(qn('w:val'), 'single')  # 單線
                border.set(qn('w:sz'), '4')       # 線條大小（1/8 pt）
                border.set(qn('w:space'), '0')    # 無間距
                border.set(qn('w:color'), '000000')  # 黑色
                tc_borders.append(border)

            tc_pr.append(tc_borders)
    
    # 設置第一列（Header Row）的背景色、粗體與置中
    for cell in table.rows[0].cells:  # 只針對第一列
        shading = parse_xml(r'<w:shd {} w:fill="E2EFD9"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading)  # 設置背景顏色

        paragraph = cell.paragraphs[0]
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        run.font.bold = True  # 設置為粗體
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 讓標題置中

    # 設置第一欄（Column 1）內容置中，第二欄（Column 2）內容靠左
    for row in table.rows[1:]:  # 跳過第一列，從第二列開始
        first_col = row.cells[0]  # 第一欄
        second_col = row.cells[1]  # 第二欄

        # 第一欄置中
        first_paragraph = first_col.paragraphs[0]
        first_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 第二欄靠左
        second_paragraph = second_col.paragraphs[0]
        second_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 設置表格內字體
    for row in table.rows:
        for cell in row.cells:
            paragraph = cell.paragraphs[0]
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()

            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(0)
            paragraph_format.space_after = Pt(0)
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

            # 設定字型
            font = run.font
            font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "標楷體")
            font.size = Pt(12)
    
     # 設定欄位寬度（第一欄30%，第二欄70%）
    table.autofit = False  # 取消自動調整大小，手動設定寬度
    for row in table.rows:
        row.cells[0].width = Inches(2)  # 第一欄 30%
        row.cells[1].width = Inches(5)  # 第二欄 70%

