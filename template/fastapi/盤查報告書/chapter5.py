from docx import Document
from docx.shared import Cm
import requests

from .ch0Def import set_heading, set_heading2, set_paragraph, set_explain
from .ch5Def import set_ch5_table1, set_ch5_table2, set_ch5_table3, set_ch5_table4
from .storeDef2 import get_org_name
from .storeDef3 import get_latest_baseline_year,get_latest_cfv_start_date

def create_chapter5(user_id):
    org_name = get_org_name(user_id)
    
    # 獲取日期並轉換格式
    date_str = get_latest_cfv_start_date()
    # 假設日期格式為 YYYY-MM-DD
    date_parts = date_str.split('-')
    if len(date_parts) >= 2:
        date = f"{date_parts[0]}年{int(date_parts[1])}月"
    else:
        date = date_str  # 如果日期格式不是預期的，保留原格式
    
    year = get_latest_baseline_year()
    year2 = year-1911
    doc = Document()

    # 獲取文檔的第一個 section（默認只有一個）
    section = doc.sections[0]

    # 設置自訂邊界，單位是厘米 (cm)
    section.top_margin = Cm(2)  # 上邊距
    section.bottom_margin = Cm(2)  # 下邊距
    section.left_margin = Cm(2)  # 左邊距
    section.right_margin = Cm(2)  # 右邊距
    
    # 取得API資料
    try:
        response = requests.get("http://localhost:8000/result")  # 請替換為實際API位址
        result_data = response.json()
        electricity_usage = result_data["result"]["Electricity_Usage"]
        quantitative_inventory = result_data["result"]["Quantitative_Inventory"]
    except Exception as e:
        print(f"無法獲取API資料: {e}")
        electricity_usage = None
        quantitative_inventory = {}

    ###############第五章######################
    #標題
    title = doc.add_heading("第五章、基準年",level=1)
    set_heading(title)

    #5.1 基準年設定
    preface = doc.add_heading("5.1 基準年設定",level=2)
    set_heading2(preface)

    # 使用從API獲取的總排放當量數據
    total_emission = quantitative_inventory.get("total_emission_equivalent", "xxxx.xxxx")
    content = doc.add_paragraph(f"本機構於{date}規劃並導入溫室氣體盤查，以{year2}年度(最近一個完整會計年度)為本機構溫室氣體盤查之基準年。基準年排放清冊如表5.1所示，基準年排放量為{total_emission}噸CO2e。")
    set_paragraph(content)

    explain = doc.add_paragraph(f"表5.1、{org_name}基準年溫室氣體排放清冊")
    set_explain(explain)

    table1 = doc.add_table(rows=3, cols=12)
    set_ch5_table1(table1, electricity_usage)

    doc.add_paragraph("")  # 插入一個空白行

    table2 = doc.add_table(rows=5, cols=10)
    # Pass the quantitative_inventory to the set_ch5_table2 function
    set_ch5_table2(table2, quantitative_inventory)

    doc.add_paragraph("")  # 插入一個空白行

    table3 = doc.add_table(rows=4, cols=10)
    set_ch5_table3(table3, quantitative_inventory)

    doc.add_paragraph("")  # 插入一個空白行

    table4 = doc.add_table(rows=8, cols=7)
    set_ch5_table4(table4, quantitative_inventory)

    # 插入分頁符號
    doc.add_page_break()

    return doc