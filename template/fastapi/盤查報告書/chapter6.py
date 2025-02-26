from docx import Document
from docx.shared import Cm

from storeDef import set_heading, set_heading2, set_paragraph, set_ch6_stairs1


def create_chapter6():
    doc = Document()

        # 獲取文檔的第一個 section（默認只有一個）
    section = doc.sections[0]

    # 設置自訂邊界，單位是厘米 (cm)
    section.top_margin = Cm(2)  # 上邊距
    section.bottom_margin = Cm(2)  # 下邊距
    section.left_margin = Cm(2)  # 左邊距
    section.right_margin = Cm(2)  # 右邊距
        
    ###############第六章######################
    #標題
    title = doc.add_heading("第六章、參考文獻",level=1)
    set_heading(title)

    content = doc.add_paragraph("本報告書係參考下列文獻製作：")
    set_ch6_stairs1(content)
    content = doc.add_paragraph("1. Intergovernmental Panel on Climate Change, IPCC Guidelines for National Greenhouse Gas Inventories, 2006.10.：")
    set_ch6_stairs1(content)
    content = doc.add_paragraph("2. Intergovernmental Panel on Climate Change, The Fifth Assessment Report : Climate Change 2014 – The Science of Climate Change, 2014.")
    set_ch6_stairs1(content)
    content = doc.add_paragraph("3. ISO 14064-1：2018, Greenhouse gases - Part 1: Specification with guidance at the organization level for quantification and reporting of greenhouse gas emissions and removals.")
    set_ch6_stairs1(content)
    content = doc.add_paragraph("4. 經濟部標準檢驗局「CNS 14064溫室氣體第一部份：組織層級溫室氣體排放與移除之量化及報告附指引之規範」，中文版，110年01月。")
    set_ch6_stairs1(content)
    content = doc.add_paragraph("5. 行政院環境部「113溫室氣體盤查與登錄指引」，113年03月")
    set_ch6_stairs1(content)
   

    # 插入分頁符號
    doc.add_page_break()

    title = doc.add_heading("附件",level=1)
    set_heading(title)


    return doc
