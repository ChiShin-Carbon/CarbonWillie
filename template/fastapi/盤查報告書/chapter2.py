from docx import Document
from docx.shared import Cm

from .storeDef2 import get_org_name,get_org_address,get_business_id

from .ch0Def import set_heading, set_heading2, set_paragraph, set_explain
from .ch2Def import set_ch2_table1



def create_chapter2(user_id):
    org_name = get_org_name(user_id)
    org_address = get_org_address(user_id)
    business_id = get_business_id(user_id)

    

    doc = Document()

        # 獲取文檔的第一個 section（默認只有一個）
    section = doc.sections[0]

    # 設置自訂邊界，單位是厘米 (cm)
    section.top_margin = Cm(2)  # 上邊距
    section.bottom_margin = Cm(2)  # 下邊距
    section.left_margin = Cm(2)  # 左邊距
    section.right_margin = Cm(2)  # 右邊距
    
    ################第二章######################
    #標題
    title = doc.add_heading("第二章、盤查邊界設定",level=1)
    set_heading(title)
    
    #2.1 組織邊界設定
    preface = doc.add_heading("2.1 組織邊界設定",level=2)
    set_heading2(preface)
    
    content = doc.add_paragraph(f"本次溫室氣體盤查專案，其組織邊界設定乃是參考ISO/CNS 14064-1:2018年版與環境部113年溫室氣體盤查指引之建議，規劃並執行符合相關設定，包括(1)控制權、(2)持有股權比例、(3)財務邊界、(4)生產配股，以及(5)在法律合約定義的特定安排下，可使用不同的整合方法論等各項規定。設定上，以{org_name}位於{org_address}的【機構盤查邊界範圍】為組織邊界，統一編號為{business_id}。")
    set_paragraph(content)

    content = doc.add_paragraph("【請插入組織邊界圖－須以紅線框出明確之邊界區域】")
    set_paragraph(content)

    photo_explain = doc.add_paragraph(f"圖二、{org_name} 組織邊界")
    set_explain(photo_explain)


    #2.2 報告邊界
    preface = doc.add_heading("2.2 報告邊界",level=2)
    set_heading2(preface)
    
    content = doc.add_paragraph("本機構報告邊界包含組織邊界的【機構盤查邊界範圍】，盤查內容包含直接排放（類別1）與能源間接排放（類別2），表2為報告邊界與排放源彙整表。")
    set_paragraph(content)

    table_explain = doc.add_paragraph(f"表2、{org_name} 報告邊界與活動源彙整表")
    set_explain(table_explain)
    # 新增表格
    table = doc.add_table(rows=3, cols=2)

    # 設定表頭
    table.cell(0, 0).text = "報告邊界"
    table.cell(0, 1).text = "排放源"


    # 第一列 (類別1)
    table.cell(1, 0).text = "直接排放源\n（類別1）"
    table.cell(1, 1).text = "1. 固定：緊急發電機-柴油\n2. 人為逸散：化糞池(CH4)\n3. 人為逸散：消防設施(滅火器)、冰水主機、飲水機、冷氣機"

    # 第二列 (類別2)
    table.cell(2, 0).text = "能源間接排放源\n（類別2）"
    table.cell(2, 1).text = "1. 台電電力\n(電號：nn-nn-nnnn-nn-n)"


    set_ch2_table1(table)



    # 插入分頁符號
    doc.add_page_break()

    return doc
